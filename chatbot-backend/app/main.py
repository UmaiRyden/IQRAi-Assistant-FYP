import os
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse
from dotenv import load_dotenv
from typing import AsyncGenerator, Optional, Dict, List
import json
import uuid
from datetime import datetime
import shutil
import tempfile
import os

# Optional dependency: python-docx (used for OBE DOCX exports).
# In serverless environments where it might not be installed, we disable DOCX export
# instead of crashing the entire app.
try:
    from docx import Document
    from docx.shared import Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    Document = None
    Inches = None
    RGBColor = None
    WD_ALIGN_PARAGRAPH = None
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not installed; DOCX export endpoints will be disabled.")

from app.models.schemas import ChatRequest, ChatResponse, HealthResponse
from app.services.agent_service import AgentService
from app.services.obe_service import OBEService
from app.services.learning_service import LearningService
from app.services.email_service import EmailService
from app.services.analytics_service import AnalyticsService
from app.services.course_advisor_service import CourseAdvisorService
from app.services.academic_advisor_service import AcademicAdvisorService
from app.db.session_store import SessionStore
from app.utils.document_processor import DocumentProcessor
from app.routes import gpa_calculator

# Load environment variables
load_dotenv()

# Initialize FastAPI app
app = FastAPI(
    title="IQRAi - Iqra University AI Assistant",
    description="FastAPI backend for Iqra University chatbot",
    version="1.0.0"
)

# Include routers
app.include_router(gpa_calculator.router)

# Configure CORS - Support both local development and production
allowed_origins = [
    # Local development frontends
    "http://localhost:3000",
    "http://localhost:3001",
    "http://127.0.0.1:3000",
    "http://127.0.0.1:3001",
]

# Add frontend URLs from environment variable(s) if provided.
# FRONTEND_URL can be a single URL or a comma-separated list, e.g.:
# FRONTEND_URL=https://iqrai-frontend.vercel.app,https://staging-iqrai-frontend.vercel.app
frontend_env = os.getenv("FRONTEND_URL")
if frontend_env:
    for origin in [o.strip() for o in frontend_env.split(",") if o.strip()]:
        if origin not in allowed_origins:
            allowed_origins.append(origin)
    print(f"✅ CORS: Added frontend URL(s) from FRONTEND_URL: {frontend_env}")
else:
    print("ℹ️ CORS: FRONTEND_URL not set, using localhost origins only")

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["*"],
    expose_headers=["*"],
)

# Initialize services
print("🔄 Initializing services...")
db_store = SessionStore()
print("✅ SessionStore initialized")

print("🔄 Initializing AgentService...")
agent_service = AgentService()
print("✅ AgentService initialized")

print("🔄 Initializing OBEService...")
obe_service = OBEService()
print("✅ OBEService initialized")

print("🔄 Initializing LearningService...")
learning_service = LearningService()
print("✅ LearningService initialized")

print("🔄 Initializing AnalyticsService...")
analytics_service = AnalyticsService()
print("✅ AnalyticsService initialized")

# Set analytics service for GPA calculator (after initialization)
from app.routes import gpa_calculator as gpa_calculator_module
gpa_calculator_module.set_analytics_service(analytics_service)

# Initialize Course Advisor Service with error handling
try:
    course_advisor_service = CourseAdvisorService()
    print(f"✅ Course Advisor Service initialized with {len(course_advisor_service.course_data)} courses")
except Exception as e:
    print(f"❌ ERROR: Failed to initialize Course Advisor Service: {str(e)}")
    import traceback
    traceback.print_exc()
    # Create a placeholder service to prevent crashes
    course_advisor_service = None

# Initialize Academic Advisor Service
try:
    academic_advisor_service = AcademicAdvisorService()
    print(f"✅ Academic Advisor Service initialized")
except Exception as e:
    print(f"❌ ERROR: Failed to initialize Academic Advisor Service: {str(e)}")
    import traceback
    traceback.print_exc()
    academic_advisor_service = None

# Initialize email service (optional - will fail gracefully if env vars not set)
email_service = None
try:
    email_service = EmailService()
    print("✅ Email service initialized")
except Exception as e:
    print(f"⚠️ Email service not initialized: {e}")
    print("   Set SMTP_USER, SMTP_PASS, and SYSTEM_EMAIL environment variables to enable email functionality.")

# Note: Whisper transcription is now handled in the frontend via OpenAI Whisper API
# No local Whisper model needed


@app.on_event("startup")
async def startup_event():
    """Initialize database and services on startup."""
    await db_store.initialize()
    print("✅ IQRAi FastAPI Backend Started")
    print(f"📡 Server running on http://localhost:{os.getenv('BACKEND_PORT', 8000)}")
    print("🔗 CORS enabled for http://localhost:3000")


@app.on_event("shutdown")
async def shutdown_event():
    """Cleanup on shutdown."""
    await db_store.close()
    print("👋 IQRAi Backend Shutting Down")


@app.get("/health", response_model=HealthResponse)
async def health_check():
    """Health check endpoint."""
    return {"status": "ok"}


@app.post("/chat/stream")
async def chat_stream(request: ChatRequest):
    """
    Stream chat responses using Server-Sent Events (SSE).
    
    Args:
        request: ChatRequest containing session_id (optional) and message
        
    Returns:
        StreamingResponse with SSE chunks
    """
    # Generate or use existing session ID
    session_id = request.session_id or str(uuid.uuid4())
    
    # Get or create session history
    history = await db_store.get_session_history(session_id)
    
    # Save user message
    await db_store.save_message(
        session_id=session_id,
        role="user",
        content=request.message
    )
    
    async def event_generator() -> AsyncGenerator[str, None]:
        """Generate SSE events from agent responses."""
        try:
            # Send session ID first
            yield f"data: {json.dumps({'type': 'session', 'session_id': session_id})}\n\n"
            
            # Stream agent response
            full_response = ""
            async for chunk in agent_service.stream_chat_response(
                session_id=session_id,
                message=request.message,
                history=history
            ):
                full_response += chunk
                # Send each chunk as SSE
                yield f"data: {json.dumps({'type': 'token', 'content': chunk})}\n\n"
            
            # Save assistant response
            await db_store.save_message(
                session_id=session_id,
                role="assistant",
                content=full_response
            )
            
            # Send completion event
            yield f"data: {json.dumps({'type': 'done'})}\n\n"
            
            # Track analytics - only count successful completions
            analytics_service.track_event(
                user=session_id,
                action="ai_chat",
                meta={"message_length": len(request.message), "response_length": len(full_response)}
            )
            
        except Exception as e:
            error_msg = f"Error: {str(e)}"
            yield f"data: {json.dumps({'type': 'error', 'content': error_msg})}\n\n"
    
    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"
        }
    )


@app.get("/sessions/{session_id}/history")
async def get_session_history(session_id: str):
    """
    Retrieve chat history for a specific session.
    
    Args:
        session_id: The session identifier
        
    Returns:
        List of messages in the session
    """
    history = await db_store.get_session_history(session_id)
    return {"session_id": session_id, "history": history}


@app.delete("/sessions/{session_id}")
async def delete_session(session_id: str):
    """
    Delete a session and its history.
    
    Args:
        session_id: The session identifier
        
    Returns:
        Success message
    """
    await db_store.delete_session(session_id)
    # Also cleanup Pinecone namespace if exists
    await agent_service.cleanup_session(session_id)
    return {"message": f"Session {session_id} deleted successfully"}


@app.post("/upload/document")
async def upload_document(
    file: UploadFile = File(...),
    session_id: Optional[str] = Form(None)
):
    """
    Upload and process a document (PDF, DOCX, TXT).
    
    Args:
        file: The uploaded file
        session_id: Optional session ID for the document
        
    Returns:
        Processing results with session_id and statistics
    """
    # Generate session ID if not provided
    if not session_id:
        session_id = str(uuid.uuid4())
    
    # Validate file type
    if not DocumentProcessor.validate_file_type(file.filename):
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type. Supported types: PDF, DOCX, TXT"
        )
    
    # Create upload directory
    upload_dir = "uploaded_docs"
    os.makedirs(upload_dir, exist_ok=True)
    
    # Generate unique filename
    file_extension = os.path.splitext(file.filename)[1]
    unique_filename = f"{session_id}_{uuid.uuid4().hex[:8]}{file_extension}"
    file_path = os.path.join(upload_dir, unique_filename)
    
    try:
        # Save uploaded file
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Get file info
        file_info = DocumentProcessor.get_file_info(file_path)
        
        # Check file size (max 10MB)
        if file_info["size_mb"] > 10:
            os.remove(file_path)
            raise HTTPException(
                status_code=400,
                detail="File size too large. Maximum size is 10MB"
            )
        
        # Process and embed document
        try:
            processing_stats = await agent_service.embed_and_upsert(file_path, session_id)
            
            # Track analytics
            analytics_service.track_event(
                user=session_id,
                action="document_upload",
                meta={
                    "filename": file.filename,
                    "file_size_mb": file_info["size_mb"],
                    "file_type": file_extension
                }
            )
            
            return {
                "success": True,
                "session_id": session_id,
                "filename": file.filename,
                "file_size_mb": file_info["size_mb"],
                "chunks_processed": processing_stats["chunks_processed"],
                "total_characters": processing_stats["total_characters"],
                "message": f"Document '{file.filename}' uploaded and processed successfully!"
            }
        except ValueError as e:
            # Clean up file on processing error
            if os.path.exists(file_path):
                os.remove(file_path)
            raise HTTPException(status_code=400, detail=str(e))
        
    except HTTPException:
        raise
    except Exception as e:
        # Clean up file on any error
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(
            status_code=500,
            detail=f"Error processing document: {str(e)}"
        )


@app.get("/sessions/{session_id}/documents")
async def get_session_documents(session_id: str):
    """
    Get list of documents uploaded in a session.
    
    Args:
        session_id: The session identifier
        
    Returns:
        List of documents with metadata
    """
    upload_dir = "uploaded_docs"
    documents = []
    
    if os.path.exists(upload_dir):
        for filename in os.listdir(upload_dir):
            if filename.startswith(session_id):
                file_path = os.path.join(upload_dir, filename)
                try:
                    file_info = DocumentProcessor.get_file_info(file_path)
                    # Extract original filename (remove session_id prefix and uuid)
                    original_name = "_".join(filename.split("_")[2:]) if "_" in filename else filename
                    documents.append({
                        "filename": original_name,
                        "size_mb": file_info["size_mb"],
                        "uploaded_at": datetime.fromtimestamp(
                            os.path.getctime(file_path)
                        ).isoformat()
                    })
                except Exception as e:
                    print(f"Error reading file info: {e}")
                    continue
    
    return {
        "session_id": session_id,
        "documents": documents,
        "count": len(documents)
    }


# ============================================================================
# OBE VERIFICATION ENDPOINTS
# ============================================================================

@app.get("/obe/domains")
async def get_obe_domains():
    """
    Get available OBE domains and their levels.
    
    Returns:
        Dictionary with domain information
    """
    try:
        domains = obe_service.get_available_domains()
        domain_info = obe_service.get_domain_info()
        
        return {
            "success": True,
            "domains": domains,
            "domain_info": domain_info
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/obe/verify")
async def verify_obe_document(
    file: UploadFile = File(...),
    domain: str = Form(...),
    level: str = Form(None)  # Optional - if not provided, auto-detect levels
):
    """
    Verify an exam paper against OBE (Bloom's Taxonomy) requirements.
    
    Args:
        file: The exam paper file (PDF, DOCX, TXT)
        domain: Domain to verify (Cognitive, Psychomotor, Affective)
        level: Level to verify (e.g., C1-C6, P1-P4, A1-A4)
        
    Returns:
        Verification results with matched verbs and analysis
    """
    # Validate file type
    if not DocumentProcessor.validate_file_type(file.filename):
        raise HTTPException(
            status_code=400,
            detail="Unsupported file type. Supported types: PDF, DOCX, TXT"
        )
    
    # Create temporary directory for OBE documents
    obe_dir = "obe_uploads"
    os.makedirs(obe_dir, exist_ok=True)
    
    # Generate unique filename
    file_extension = os.path.splitext(file.filename)[1]
    unique_filename = f"obe_{uuid.uuid4().hex[:12]}{file_extension}"
    file_path = os.path.join(obe_dir, unique_filename)
    
    try:
        # Save uploaded file
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Get file info
        file_info = DocumentProcessor.get_file_info(file_path)
        
        # Check file size (max 10MB)
        if file_info["size_mb"] > 10:
            os.remove(file_path)
            raise HTTPException(
                status_code=400,
                detail="File size too large. Maximum size is 10MB"
            )
        
        # Verify document using OBE service
        try:
            # If level not provided, use auto-detection
            if not level:
                verification_result = await obe_service.verify_document_auto_detect(
                    file_path=file_path,
                    domain=domain
                )
            else:
                # Legacy mode: verify against specific level
                verification_result = await obe_service.verify_document(
                    file_path=file_path,
                    domain=domain,
                    level=level
                )
            
            # Add file info to result
            verification_result["filename"] = file.filename
            verification_result["file_size_mb"] = file_info["size_mb"]
            verification_result["verification_timestamp"] = datetime.now().isoformat()
            
            # Track analytics
            analytics_service.track_event(
                user="system",  # Could use session_id if provided
                action="obe_verification",
                meta={
                    "domain": domain,
                    "level": level or "auto-detect",
                    "total_questions": verification_result.get("total_questions", 0),
                    "match_percentage": verification_result.get("match_percentage", 0),
                    "status": verification_result.get("status", "Unknown")
                }
            )
            
            return {
                "success": True,
                **verification_result
            }
            
        except ValueError as e:
            # Clean up file on verification error
            if os.path.exists(file_path):
                os.remove(file_path)
            raise HTTPException(status_code=400, detail=str(e))
        finally:
            # Clean up the file after verification
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except Exception as e:
                    print(f"Warning: Failed to delete temporary file: {e}")
        
    except HTTPException:
        raise
    except Exception as e:
        # Clean up file on any error
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
            except:
                pass
        raise HTTPException(
            status_code=500,
            detail=f"Error processing document: {str(e)}"
        )


@app.get("/obe/verbs/{domain}/{level}")
async def get_obe_verbs(domain: str, level: str):
    """
    Get the list of required verbs for a specific domain and level.
    
    Args:
        domain: Domain (Cognitive, Psychomotor, Affective)
        level: Level (e.g., C1-C6, P1-P4, A1-A4)
        
    Returns:
        List of required verbs
    """
    try:
        verbs = obe_service.get_verbs_for_level(domain, level)
        
        if not verbs:
            raise HTTPException(
                status_code=404,
                detail=f"No verbs found for domain '{domain}' and level '{level}'"
            )
        
        return {
            "success": True,
            "domain": domain,
            "level": level,
            "verbs": verbs,
            "count": len(verbs)
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/obe/rewrite-question")
async def rewrite_obe_question(
    question: str = Form(...),
    domain: str = Form(...),
    level: str = Form(...),
    context: Optional[str] = Form(None)
):
    """
    Rewrite a question to match a specific Bloom's Taxonomy level.
    
    Args:
        question: Original question text
        domain: Domain (Cognitive, Psychomotor, Affective)
        level: Target level (C1-C6, P1-P4, A1-A4)
        context: Optional additional context for the question
        
    Returns:
        Improved question text
    """
    try:
        improved_question = await obe_service.rewrite_question(
            question=question,
            domain=domain,
            level=level,
            context=context
        )
        
        # Track analytics
        analytics_service.track_event(
            user="system",
            action="obe_rewrite",
            meta={
                "domain": domain,
                "level": level,
                "has_context": context is not None
            }
        )
        
        return {
            "success": True,
            "improved_question": improved_question,
            "original_question": question,
            "domain": domain,
            "level": level
        }
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to rewrite question: {str(e)}"
        )


@app.post("/obe/verify-single-question")
async def verify_single_question(
    question: str = Form(...),
    domain: str = Form(...)
):
    """
    Verify a single question to detect its Bloom's Taxonomy level and matched verbs.
    
    Args:
        question: Question text to verify
        domain: Domain (Cognitive, Psychomotor, Affective)
        
    Returns:
        Detection results with levels and matched verbs
    """
    try:
        # Use the OBE service's detection method
        detected_levels = obe_service._detect_level_for_question(question, domain)
        
        # Determine status - check if dict has any keys
        has_match = detected_levels and len(detected_levels) > 0
        status = 'Match Found' if has_match else 'No CLO-level match found'
        
        # Convert empty dict to null for consistency with frontend
        detected_levels_result = detected_levels if has_match else None
        
        return {
            "success": True,
            "question": question,
            "domain": domain,
            "detected_levels": detected_levels_result,
            "status": status
        }
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to verify question: {str(e)}"
        )


def create_docx_from_obe_questions(questions: List[Dict], filename: str = "obe_questions.docx") -> str:
    """
    Create a DOCX file from OBE verification questions.
    
    Args:
        questions: List of question dictionaries with question_number and question_text
        filename: Output filename
        
    Returns:
        Path to created DOCX file
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed; DOCX export is unavailable.")

    doc = Document()
    
    # Add title
    title = doc.add_heading("OBE VERIFIED QUESTIONS", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Add metadata
    doc.add_paragraph(f"Total Questions: {len(questions)}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("")  # Empty line
    
    # Add questions
    for q in questions:
        q_num = q.get('question_number', '?')
        q_text = q.get('question_text', '')
        doc.add_heading(f"Question {q_num}", level=2)
        doc.add_paragraph(q_text)
        doc.add_paragraph("")  # Empty line between questions
    
    # Save document
    obe_dir = "obe_uploads"
    os.makedirs(obe_dir, exist_ok=True)
    output_path = os.path.join(obe_dir, filename)
    doc.save(output_path)
    return output_path


@app.post("/obe/download-questions")
async def download_obe_questions(
    questions: str = Form(...),  # JSON string of questions array
    format: str = Form("docx"),  # "docx" or "txt"
    filename: Optional[str] = Form(None)
):
    """
    Download OBE verified questions as DOCX or TXT file.
    
    Args:
        questions: JSON string containing questions array
        format: File format ("docx" or "txt")
        filename: Optional custom filename
        
    Returns:
        File download
    """
    if not DOCX_AVAILABLE and format.lower() == "docx":
        raise HTTPException(
            status_code=500,
            detail="DOCX export is not available on this deployment (python-docx not installed)."
        )

    try:
        # Parse questions
        questions_list = json.loads(questions)
        
        if format.lower() == "docx":
            # Generate filename if not provided
            if not filename:
                filename = f"obe_questions_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            
            # Create DOCX file
            docx_path = create_docx_from_obe_questions(questions_list, filename)
            
            # Return file for download
            return FileResponse(
                path=docx_path,
                filename=filename,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:  # TXT format
            # Generate filename if not provided
            if not filename:
                filename = f"obe_questions_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
            
            # Create TXT content
            content = "OBE VERIFIED QUESTIONS\n"
            content += f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
            
            for q in questions_list:
                q_num = q.get('question_number', '?')
                q_text = q.get('question_text', '')
                content += f"Question {q_num}\n"
                content += f"{q_text}\n\n"
            
            # Save to temporary file
            import tempfile
            with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
                f.write(content)
                temp_path = f.name
            
            return FileResponse(
                path=temp_path,
                filename=filename,
                media_type="text/plain"
            )
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid JSON format for questions")
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to create download file: {str(e)}"
        )


@app.post("/obe/generate")
async def generate_obe_questions(
    domain: str = Form(...),
    levels: str = Form(...),  # JSON array string like '["C1", "C2"]'
    count: int = Form(...),
    lecture_text: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None)
):
    """
    Generate OBE-aligned questions from lecture content.
    
    Args:
        domain: Domain (Cognitive, Psychomotor, Affective)
        levels: JSON string array of Bloom levels (e.g., '["C1", "C2", "C3"]')
        count: Number of questions to generate
        lecture_text: Optional text content (if provided, use this)
        file: Optional file upload (PDF, DOCX, TXT, PPTX) - extract text if provided
        
    Returns:
        List of generated questions with their Bloom levels
    """
    try:
        # Validate domain
        if domain not in ['Cognitive', 'Psychomotor', 'Affective']:
            raise HTTPException(status_code=400, detail=f"Invalid domain: {domain}")
        
        # Parse levels
        try:
            levels_list = json.loads(levels)
            if not isinstance(levels_list, list) or len(levels_list) == 0:
                raise ValueError("Levels must be a non-empty array")
        except (json.JSONDecodeError, ValueError) as e:
            raise HTTPException(status_code=400, detail=f"Invalid levels format: {str(e)}")
        
        # Validate count
        if count < 1 or count > 20:
            raise HTTPException(status_code=400, detail="Count must be between 1 and 20")
        
        # Get lecture text - either from file or provided text
        content_text = ""
        
        if file:
            # Extract text from uploaded file
            if not DocumentProcessor.validate_file_type(file.filename):
                raise HTTPException(
                    status_code=400,
                    detail="Unsupported file type. Supported: PDF, DOCX, DOC, TXT"
                )
            
            # Save temporarily
            obe_dir = "obe_uploads"
            os.makedirs(obe_dir, exist_ok=True)
            file_path = os.path.join(obe_dir, f"temp_{uuid.uuid4()}{os.path.splitext(file.filename)[1]}")
            
            try:
                with open(file_path, "wb") as buffer:
                    shutil.copyfileobj(file.file, buffer)
                
                # Extract text
                content_text = DocumentProcessor.extract_text(file_path)
                
                # Clean up
                if os.path.exists(file_path):
                    os.remove(file_path)
                    
            except Exception as e:
                if os.path.exists(file_path):
                    try:
                        os.remove(file_path)
                    except:
                        pass
                raise HTTPException(status_code=400, detail=f"Failed to extract text from file: {str(e)}")
        elif lecture_text:
            content_text = lecture_text.strip()
        else:
            raise HTTPException(status_code=400, detail="Either file or lecture_text must be provided")
        
        if not content_text or len(content_text) < 50:
            raise HTTPException(status_code=400, detail="Lecture content too short. Provide at least 50 characters.")
        
        # Generate questions
        questions = await obe_service.generate_questions(
            lecture_text=content_text,
            domain=domain,
            levels=levels_list,
            count=count
        )
        
        if not questions:
            raise HTTPException(status_code=500, detail="Failed to generate questions. Please try again.")
        
        return {
            "success": True,
            "domain": domain,
            "generated_count": len(questions),
            "questions": questions
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate questions: {str(e)}"
        )


# ============================================================================
# LEARNING ENDPOINTS
# ============================================================================

# In-memory storage for learning materials (replace with DB in production)
learning_materials = {}  # session_id -> {files: [], content: str}
learning_history = {}  # session_id -> {quizzes: [], socratic: [], study: []}

# Global directory for learning uploads
LEARNING_DIR = "learning_uploads"


@app.post("/learning/upload")
async def upload_learning_material(
    file: UploadFile = File(...),
    session_id: Optional[str] = Form(None)
):
    """
    Upload course material for learning features.
    
    Args:
        file: Course material (PDF, DOCX, PPTX, TXT)
        session_id: Optional session ID
        
    Returns:
        Upload confirmation with extracted content info
    """
    # Validate file type
    allowed_types = [".pdf", ".docx", ".doc", ".txt", ".pptx"]
    file_ext = "." + file.filename.split(".")[-1].lower()
    
    if file_ext not in allowed_types:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type. Allowed: {', '.join(allowed_types)}"
        )
    
    # Generate session if not provided
    if not session_id:
        session_id = str(uuid.uuid4())
    
    # Create learning directory
    os.makedirs(LEARNING_DIR, exist_ok=True)
    
    # Save file
    file_id = str(uuid.uuid4())
    file_path = os.path.join(LEARNING_DIR, f"{file_id}_{file.filename}")
    
    try:
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Extract text
        try:
            content = DocumentProcessor.extract_text(file_path)
        except Exception as e:
            # For PPTX or unsupported, return partial support message
            content = f"File uploaded but text extraction not fully supported for {file_ext}. Using filename: {file.filename}"
        
        # Store in memory
        if session_id not in learning_materials:
            learning_materials[session_id] = {"files": [], "content": ""}
        
        learning_materials[session_id]["files"].append({
            "file_id": file_id,
            "filename": file.filename,
            "path": file_path,
            "size_mb": os.path.getsize(file_path) / (1024 * 1024),
            "content": content  # Store content per file
        })
        learning_materials[session_id]["content"] += "\n\n" + content
        
        return {
            "success": True,
            "session_id": session_id,
            "file_id": file_id,
            "filename": file.filename,
            "file_size_mb": round(os.path.getsize(file_path) / (1024 * 1024), 2),
            "content_preview": content[:300] + "..." if len(content) > 300 else content,
            "total_files": len(learning_materials[session_id]["files"])
        }
        
    except Exception as e:
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/learning/materials/{session_id}/{file_id}")
async def delete_learning_material(session_id: str, file_id: str):
    """
    Delete an uploaded learning material file.
    
    Args:
        session_id: Session ID
        file_id: File ID to delete
        
    Returns:
        Deletion confirmation
    """
    if session_id not in learning_materials:
        raise HTTPException(status_code=404, detail="Session not found")
    
    # Find and remove file
    file_to_delete = None
    for file in learning_materials[session_id]["files"]:
        if file["file_id"] == file_id:
            file_to_delete = file
            break
    
    if not file_to_delete:
        raise HTTPException(status_code=404, detail="File not found")
    
    # Delete physical file
    if os.path.exists(file_to_delete["path"]):
        try:
            os.remove(file_to_delete["path"])
        except Exception as e:
            print(f"Warning: Could not delete file {file_to_delete['path']}: {e}")
    
    # Remove from memory
    learning_materials[session_id]["files"] = [
        f for f in learning_materials[session_id]["files"] if f["file_id"] != file_id
    ]
    
    # Rebuild content without deleted file
    learning_materials[session_id]["content"] = "\n\n".join(
        f["content"] for f in learning_materials[session_id]["files"]
    )
    
    return {
        "success": True,
        "message": f"File {file_to_delete['filename']} deleted successfully"
    }


@app.post("/learning/generate_quiz")
async def generate_quiz(
    session_id: str = Form(...),
    quiz_type: str = Form(...),  # mcq, true_false, short_answer
    num_questions: str = Form("5"),  # Accept as string, convert to int
    total_marks: Optional[str] = Form(None),  # Optional: accept as string, convert to int
    content_type: Optional[str] = Form(None),  # Optional: quiz or assignment
    file_id: Optional[str] = Form(None)  # Optional: specific file to use
):
    """
    Generate a quiz based on uploaded content with marks-based complexity.
    
    Args:
        session_id: Session with uploaded materials
        quiz_type: Type of quiz (mcq, true_false, short_answer)
        num_questions: Number of questions (default 5)
        total_marks: Total marks for the quiz/assignment (default 10)
        content_type: "quiz" or "assignment" (default "quiz")
        file_id: Optional file ID to use specific file
        
    Returns:
        Generated quiz JSON
    """
    if session_id not in learning_materials:
        raise HTTPException(status_code=404, detail="No materials found. Please upload content first.")
    
    # Convert string parameters to proper types
    try:
        num_questions_int = int(num_questions) if num_questions else 5
    except (ValueError, TypeError):
        num_questions_int = 5
    
    try:
        total_marks_int = int(total_marks) if total_marks else 10
    except (ValueError, TypeError):
        total_marks_int = 10
    
    content_type_str = content_type if content_type else "quiz"
    
    # Use specific file or all content
    content = learning_materials[session_id]["content"]
    
    if file_id:
        # Find specific file
        file = next((f for f in learning_materials[session_id]["files"] if f["file_id"] == file_id), None)
        if file:
            content = file["content"]
        else:
            raise HTTPException(status_code=404, detail="File not found")
    
    if not content or len(content.strip()) < 50:
        raise HTTPException(status_code=400, detail="Insufficient content to generate quiz")
    
    try:
        quiz = await learning_service.generate_quiz(
            content, 
            quiz_type, 
            num_questions_int,
            total_marks=total_marks_int,
            content_type=content_type_str
        )
        
        # Store in history
        if session_id not in learning_history:
            learning_history[session_id] = {"quizzes": [], "socratic": [], "study": []}
        
        quiz_record = {
            "quiz_id": str(uuid.uuid4()),
            "quiz_type": quiz_type,
            "num_questions": num_questions_int,
            "total_marks": total_marks_int,
            "content_type": content_type_str,
            "file_id": file_id,
            "quiz_data": quiz,
            "timestamp": datetime.now().isoformat()
        }
        learning_history[session_id]["quizzes"].append(quiz_record)
        
        # Track analytics
        analytics_service.track_event(
            user=session_id,
            action="quiz_generated",
            meta={
                "quiz_type": quiz_type,
                "num_questions": num_questions_int,
                "total_marks": total_marks_int,
                "content_type": content_type_str
            }
        )
        
        return {
            "success": True,
            "quiz_id": quiz_record["quiz_id"],
            "total_marks": total_marks_int,
            "marks_per_question": round(total_marks_int / num_questions_int, 1),
            **quiz
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/learning/socratic")
async def socratic_dialogue(
    session_id: str = Form(...),
    topic: str = Form(...),
    conversation_history: Optional[str] = Form(None)  # JSON string
):
    """
    Engage in Socratic dialogue.
    
    Args:
        session_id: Session ID
        topic: Topic or question from student
        conversation_history: Optional JSON string of previous messages
        
    Returns:
        Socratic response
    """
    # Parse conversation history
    history = []
    if conversation_history:
        try:
            history = json.loads(conversation_history)
        except:
            history = []
    
    try:
        response = await learning_service.socratic_dialogue(topic, history)
        
        # Store in history
        if session_id not in learning_history:
            learning_history[session_id] = {"quizzes": [], "socratic": [], "study": []}
        
        learning_history[session_id]["socratic"].append({
            "topic": topic,
            "response": response,
            "timestamp": datetime.now().isoformat()
        })
        
        return {
            "success": True,
            "response": response
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/learning/study_mode/question")
async def study_mode_question(
    session_id: str = Form(...),
    file_id: Optional[str] = Form(None)  # Optional: specific file to use
):
    """
    Get a study mode question based on uploaded content.
    
    Args:
        session_id: Session with uploaded materials
        file_id: Optional file ID to use specific file
        
    Returns:
        Study question
    """
    if session_id not in learning_materials:
        raise HTTPException(status_code=404, detail="No materials found. Please upload content first.")
    
    # Use specific file or all content
    content = learning_materials[session_id]["content"]
    
    if file_id:
        file = next((f for f in learning_materials[session_id]["files"] if f["file_id"] == file_id), None)
        if file:
            content = file["content"]
        else:
            raise HTTPException(status_code=404, detail="File not found")
    
    # Get previous Q&A
    previous_qa = []
    if session_id in learning_history and "study" in learning_history[session_id]:
        previous_qa = learning_history[session_id]["study"]
    
    try:
        question = await learning_service.study_mode_question(content, previous_qa)
        
        return {
            "success": True,
            "question": question
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/learning/study_mode/answer")
async def study_mode_answer(
    session_id: str = Form(...),
    question: str = Form(...),
    answer: str = Form(...)
):
    """
    Submit answer for study mode question and get feedback.
    
    Args:
        session_id: Session ID
        question: The question asked
        answer: Student's answer
        
    Returns:
        Feedback on answer
    """
    if session_id not in learning_materials:
        raise HTTPException(status_code=404, detail="No materials found.")
    
    content = learning_materials[session_id]["content"]
    
    try:
        feedback = await learning_service.evaluate_study_answer(question, answer, content)
        
        # Store in history
        if session_id not in learning_history:
            learning_history[session_id] = {"quizzes": [], "socratic": [], "study": []}
        
        learning_history[session_id]["study"].append({
            "question": question,
            "answer": answer,
            "feedback": feedback,
            "timestamp": datetime.now().isoformat()
        })
        
        # Track analytics - count study sessions when answer is submitted
        analytics_service.track_event(
            user=session_id,
            action="study_session",
            meta={
                "has_feedback": feedback is not None
            }
        )
        
        return {
            "success": True,
            "feedback": feedback
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/learning/history/{session_id}")
async def get_learning_history(session_id: str):
    """
    Get learning history for a session.
    
    Args:
        session_id: Session ID
        
    Returns:
        History of quizzes, Socratic sessions, and study mode
    """
    if session_id not in learning_history:
        return {
            "success": True,
            "session_id": session_id,
            "quizzes": [],
            "socratic": [],
            "study": []
        }
    
    return {
        "success": True,
        "session_id": session_id,
        **learning_history[session_id]
    }


@app.get("/learning/materials/{session_id}")
async def get_learning_materials(session_id: str):
    """
    Get uploaded materials for a session.
    
    Args:
        session_id: Session ID
        
    Returns:
        List of uploaded files
    """
    if session_id not in learning_materials:
        return {
            "success": True,
            "session_id": session_id,
            "files": []
        }
    
    return {
        "success": True,
        "session_id": session_id,
        "files": learning_materials[session_id]["files"],
        "total_files": len(learning_materials[session_id]["files"])
    }


def create_docx_from_quiz(quiz_data: Dict, filename: str = "quiz.docx") -> str:
    """
    Create a DOCX file from quiz data.
    
    Args:
        quiz_data: Quiz data containing questions
        filename: Output filename
        
    Returns:
        Path to created DOCX file
    """
    doc = Document()
    
    # Add title with black color
    title = doc.add_heading(f"{quiz_data.get('quiz_type', 'Quiz').upper()}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Set title text to black
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
    
    # Add metadata
    doc.add_paragraph(f"Total Questions: {len(quiz_data.get('questions', []))}")
    doc.add_paragraph(f"Total Marks: {quiz_data.get('total_marks', 'N/A')}")
    doc.add_paragraph(f"Marks per Question: {quiz_data.get('marks_per_question', 'N/A')}")
    doc.add_paragraph("")  # Empty line
    
    # Add questions (without answers/explanations for student version)
    for i, question in enumerate(quiz_data.get('questions', []), 1):
        # Question number and text
        q_heading = doc.add_heading(f"Question {i} ({question.get('marks', 'N/A')} marks)", level=2)
        # Set question heading text to black
        for run in q_heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
        doc.add_paragraph(question.get('question', ''))
        
        # Handle different question types - only show options, not answers
        if question.get('options'):
            # MCQ options (no correct answer shown)
            for option in question.get('options', []):
                doc.add_paragraph(f"  {option}", style='List Bullet')
        
        doc.add_paragraph("")  # Empty line between questions
    
    # Save document - ensure learning directory exists
    os.makedirs(LEARNING_DIR, exist_ok=True)
    output_path = os.path.join(LEARNING_DIR, filename)
    doc.save(output_path)
    return output_path


@app.post("/learning/download_docx")
async def download_quiz_as_docx(
    quiz_data: str = Form(...),  # JSON string of quiz data
    filename: str = Form("quiz.docx")
):
    """
    Download generated quiz/assignment as DOCX file.
    
    Args:
        quiz_data: JSON string containing quiz data
        filename: Desired filename for download
        
    Returns:
        DOCX file download
    """
    try:
        # Parse quiz data
        quiz_dict = json.loads(quiz_data)
        
        # Create DOCX file
        docx_path = create_docx_from_quiz(quiz_dict, filename)
        
        # Return file for download
        return FileResponse(
            path=docx_path,
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Invalid quiz data format")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creating DOCX: {str(e)}")


# Note: Transcription is now handled in the frontend via OpenAI Whisper API
# This endpoint is disabled but kept for backwards compatibility
@app.post("/transcribe")
async def transcribe_audio(file: UploadFile = File(...)):
    """
    Transcription endpoint disabled - use frontend OpenAI Whisper API instead.
    
    This endpoint is kept for backwards compatibility but returns an error.
    Voice transcription is now handled directly in the frontend using OpenAI's Whisper API.
    """
    raise HTTPException(
        status_code=410,  # 410 Gone - resource no longer available
        detail="Transcription endpoint is disabled. Voice transcription is now handled in the frontend via OpenAI Whisper API. Please use the frontend voice input feature."
    )


@app.post("/send-email")
async def send_email(
    to: str = Form(...),
    subject: str = Form(...),
    body: str = Form(...),
    reply_to: Optional[str] = Form(None)
):
    """
    Send email via Gmail SMTP with send-on-behalf behavior.
    
    Args:
        to: Department email address (recipient)
        subject: Email subject (editable)
        body: Email body content (AI-generated application text)
        reply_to: Student email for Reply-To header (optional)
        
    Returns:
        JSON response with status
        
    Email Rules:
        - From: SYSTEM_EMAIL (system email from env)
        - Reply-To: student email (if provided)
        - To: department email (selected by user)
        - Subject: dynamic from request
        - Body: AI-generated application text
    """
    try:
        if not email_service:
            raise HTTPException(
                status_code=503,
                detail="Email service is not configured. Please set SMTP_USER, SMTP_PASS, and SYSTEM_EMAIL environment variables."
            )
        
        # Validate required fields
        if not to or not subject or not body:
            raise HTTPException(
                status_code=400,
                detail="Missing required fields: to, subject, and body are required"
            )
        
        # Validate email format (basic check)
        if "@" not in to:
            raise HTTPException(status_code=400, detail="Invalid recipient email format")
        
        # Send email
        email_service.send_email(
            to_email=to,
            subject=subject,
            body=body,
            reply_to=reply_to
        )
        
        # Track event
        analytics_service.track_event(
            user=reply_to or "anonymous",
            action="EMAIL_SENT",
            meta={"to": to, "subject": subject[:50], "type": "study_plan" if "study plan" in subject.lower() else "general"}
        )
        
        return {"status": "sent"}
        
    except HTTPException:
        raise
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        print(f"📧 Error sending email: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to send email: {str(e)}"
        )


# Statistics endpoints
@app.post("/stats/track")
async def track_event(request: dict = Body(...)):
    """
    Track a custom event from frontend.
    
    Request body:
    {
        "user": "user_id",
        "action": "voice_assistant_call",
        "meta": {...}
    }
    """
    try:
        user = request.get("user", "anonymous")
        action = request.get("action", "")
        meta = request.get("meta", {})
        
        analytics_service.track_event(
            user=user,
            action=action,
            meta=meta
        )
        
        return {"success": True}
    except Exception as e:
        print(f"Error tracking event: {e}")
        return {"success": False, "error": str(e)}


@app.get("/stats/dashboard")
async def get_dashboard_stats():
    """
    Get aggregated statistics for the dashboard.
    
    Returns:
        Dictionary with KPI counts and chart data
    """
    try:
        stats = analytics_service.get_dashboard_stats()
        return {
            "success": True,
            **stats
        }
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to retrieve dashboard stats: {str(e)}"
        )


@app.get("/stats/recent")
async def get_recent_activity(limit: int = 20):
    """
    Get recent activity events.
    
    Args:
        limit: Maximum number of events to return (default 20)
        
    Returns:
        List of recent events
    """
    try:
        events = analytics_service.get_recent_events(limit=limit)
        return {
            "success": True,
            "events": events,
            "count": len(events)
        }
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to retrieve recent activity: {str(e)}"
        )


# Course Advisor endpoints
@app.post("/course-advisor/upload-transcript")
async def upload_transcript(
    file: UploadFile = File(...),
    session_id: Optional[str] = Form(None)
):
    """
    Upload and analyze a student transcript.
    
    Args:
        file: Transcript file (PDF, DOCX, TXT)
        session_id: Optional session ID
        
    Returns:
        Analysis results with completed courses, recommendations, and FYP progress
    """
    # Check if service is initialized
    if course_advisor_service is None:
        raise HTTPException(
            status_code=500,
            detail="Course Advisor service is not available. Please check server logs."
        )
    
    # Validate file type
    if not file or not file.filename:
        raise HTTPException(
            status_code=400,
            detail="No file provided or filename missing"
        )
    
    print(f"Received file upload request: {file.filename}, Content-Type: {file.content_type}")
    
    if not DocumentProcessor.validate_file_type(file.filename):
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {os.path.splitext(file.filename)[1]}. Supported types: PDF, DOCX, TXT"
        )
    
    # Generate session ID if not provided
    if not session_id:
        session_id = str(uuid.uuid4())
    
    # Create upload directory
    upload_dir = "course_advisor_uploads"
    os.makedirs(upload_dir, exist_ok=True)
    
    # Generate unique filename
    file_extension = os.path.splitext(file.filename)[1]
    unique_filename = f"{session_id}_{uuid.uuid4().hex[:8]}{file_extension}"
    file_path = os.path.join(upload_dir, unique_filename)
    
    file_path_created = False
    
    try:
        # Save uploaded file
        print(f"Saving file to: {file_path}")
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        file_path_created = True
        print(f"File saved successfully. Size: {os.path.getsize(file_path)} bytes")
        
        # Extract text from transcript
        try:
            print(f"Extracting text from: {file_path}")
            transcript_text = DocumentProcessor.extract_text(file_path)
            print(f"Extracted text length: {len(transcript_text) if transcript_text else 0}")
        except Exception as e:
            print(f"Error extracting text: {str(e)}")
            import traceback
            traceback.print_exc()
            if file_path_created and os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except:
                    pass
            error_detail = str(e)
            # Provide helpful guidance for image-based PDFs
            if "image-based" in error_detail.lower() or "scanned" in error_detail.lower():
                error_detail = error_detail + "\n\n💡 Tip: If your transcript is a scanned image, try:\n" \
                    "- Converting it to a text-based PDF using OCR tools\n" \
                    "- Copying the text and saving it as a .txt file\n" \
                    "- Converting to .docx format"
            
            raise HTTPException(
                status_code=400,
                detail=error_detail
            )
        
        if not transcript_text or len(transcript_text.strip()) < 50:
            print(f"Transcript too short: {len(transcript_text) if transcript_text else 0} chars")
            if file_path_created and os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except:
                    pass
            raise HTTPException(
                status_code=400,
                detail=f"Transcript appears to be empty or too short ({len(transcript_text) if transcript_text else 0} chars). Please upload a valid transcript with at least 50 characters."
            )
        
        # Analyze transcript
        try:
            print(f"Starting transcript analysis...")
            analysis = await course_advisor_service.analyze_transcript(transcript_text)
            print(f"Analysis complete. Found {len(analysis.get('completed_courses', []))} completed courses.")
        except Exception as e:
            print(f"Error in transcript analysis: {str(e)}")
            import traceback
            traceback.print_exc()
            if file_path_created and os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except:
                    pass
            raise HTTPException(
                status_code=500,
                detail=f"Failed to analyze transcript: {str(e)}"
            )
        
        # Track analytics
        analytics_service.track_event(
            user=session_id,
            action="course_advisor_analysis",
            meta={
                "filename": file.filename,
                "completed_courses_count": len(analysis.get("completed_courses", [])),
                "eligible_courses_count": len(analysis.get("eligible_courses", []))
            }
        )
        
        # Clean up file (optional - you might want to keep it)
        # os.remove(file_path)
        
        return {
            "success": True,
            "session_id": session_id,
            "filename": file.filename,
            **analysis
        }
        
    except HTTPException:
        raise
    except Exception as e:
        # Clean up on error
        print(f"Unexpected error in upload_transcript: {str(e)}")
        import traceback
        traceback.print_exc()
        if 'file_path' in locals() and file_path_created and os.path.exists(file_path):
            try:
                os.remove(file_path)
            except:
                pass
        raise HTTPException(
            status_code=500,
            detail=f"Unexpected error: {str(e)}"
        )


@app.post("/course-advisor/chat")
async def course_advisor_chat(
    message: str = Form(...),
    session_id: Optional[str] = Form(None),
    context: Optional[str] = Form(None)  # JSON string with previous analysis context
):
    """
    Chat with the Course Advisor about course recommendations.
    
    Args:
        message: User's question
        session_id: Session ID
        context: Optional JSON string with previous analysis (completed courses, recommendations, etc.)
        
    Returns:
        AI response about course recommendations
    """
    try:
        # Parse context if provided
        context_data = {}
        if context:
            try:
                context_data = json.loads(context)
            except json.JSONDecodeError:
                pass
        
        # Build prompt with rich context
        prompt = message
        if context_data:
            completed = context_data.get("completed_courses", [])
            eligible = context_data.get("eligible_courses", [])
            fyp = context_data.get("fyp_progress", {})
            
            # Get detailed eligible course info
            eligible_details = []
            if eligible and course_advisor_service:
                for course_code in eligible[:15]:  # Limit to top 15 for context
                    course = next((c for c in course_advisor_service.course_data 
                                 if c.get("course_code") == course_code), None)
                    if course:
                        course_name = course.get("course_name", "")
                        is_fyp = course.get("is_fyp_prerequisite", False)
                        eligible_details.append({
                            "code": course_code,
                            "name": course_name,
                            "is_fyp_prerequisite": is_fyp
                        })
            
            # Sort: FYP prerequisites first
            eligible_details.sort(key=lambda x: (not x["is_fyp_prerequisite"], x["code"]))
            
            # Build detailed context
            eligible_list = "\n".join([
                f"- **{c['code']}** - {c['name']}" + (" (FYP Prerequisite)" if c["is_fyp_prerequisite"] else "")
                for c in eligible_details[:12]  # Top 12 courses
            ])
            
            missing_fyp = fyp.get("remaining_prereq_codes", [])[:10]
            missing_fyp_list = "\n".join([f"- {code}" for code in missing_fyp]) if missing_fyp else "None"
            
            context_str = f"""STUDENT TRANSCRIPT ANALYSIS:

**Completed Courses** ({len(completed)}):
{', '.join(completed[:20]) if completed else 'None'}{'...' if len(completed) > 20 else ''}

**Eligible Next Courses** ({len(eligible)} available):
{eligible_list if eligible_list else 'None'}

**FYP (CSC441) Progress**:
- Completed: {fyp.get('completed_prereqs', 0)}/{fyp.get('total_prereqs', 0)} prerequisites
- Remaining: {fyp.get('remaining_prereqs', 0)} prerequisites
- Missing prerequisites: {missing_fyp_list}

**Student Question**: {message}

Provide a clear, well-formatted response using the course data above. If asked for semester planning, use specific course codes and format with clear sections."""
            prompt = context_str
        
        # Import Runner and ResponseTextDeltaEvent for agent
        from agents import Runner
        from openai.types.responses import ResponseTextDeltaEvent
        
        # Run the agent
        result = Runner.run_streamed(
            course_advisor_service.agent,
            input=[{"role": "user", "content": prompt}],
        )
        
        # Stream response
        full_response = ""
        async for event in result.stream_events():
            if event.type == "raw_response_event" and isinstance(event.data, ResponseTextDeltaEvent):
                token = event.data.delta
                full_response += token
        
        return {
            "success": True,
            "response": full_response.strip()
        }
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate response: {str(e)}"
        )


# ============================================================================
# ACADEMIC ADVISOR ENDPOINTS
# ============================================================================

@app.post("/academic-advisor/generate-study-plan")
async def generate_study_plan(
    request: dict = Body(...)
):
    """
    Generate a personalized study plan for semester courses.
    
    Request body:
    {
        "courses": [{"code": "CSC101", "name": "..."}],
        "auto_save": bool
    }
    
    Returns:
        Study plan with priorities, channels, platforms, and weekly outline
    """
    if not academic_advisor_service:
        raise HTTPException(
            status_code=500,
            detail="Academic Advisor Service not initialized"
        )
    
    try:
        courses = request.get("courses", [])
        auto_save = request.get("auto_save", False)
        
        if not courses:
            raise HTTPException(
                status_code=400,
                detail="At least one course is required"
            )
        
        result = await academic_advisor_service.generate_study_plan(
            courses=courses,
            auto_save=auto_save
        )
        
        # Track event
        user_id = request.get("user_id", "anonymous") if isinstance(request, dict) else "anonymous"
        analytics_service.track_event(
            user=user_id,
            action="COURSE_ADVISOR",
            meta={"type": "study_plan", "courses_count": len(courses)}
        )
        
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate study plan: {str(e)}"
        )


@app.post("/academic-advisor/career-analysis")
async def analyze_career_path(
    transcript: Optional[UploadFile] = File(None),
    area_of_interest: str = Form(...)
):
    """
    Analyze career path and provide recommendations.
    
    Args:
        transcript: Optional transcript file (PDF or image)
        area_of_interest: Selected career area
        
    Returns:
        Career analysis with strengths, weaknesses, electives, and pathway
    """
    if not academic_advisor_service:
        raise HTTPException(
            status_code=500,
            detail="Academic Advisor Service not initialized"
        )
    
    try:
        # Validate inputs
        if not area_of_interest:
            raise HTTPException(
                status_code=400,
                detail="Area of interest is required"
            )
        
        if not transcript:
            raise HTTPException(
                status_code=400,
                detail="Transcript file is required"
            )
        
        transcript_text = None
        completed_courses = None
        
        # Process transcript
        # Save file temporarily
        file_ext = "." + transcript.filename.split(".")[-1].lower()
        allowed_exts = [".pdf", ".jpg", ".jpeg", ".png"]
        
        if file_ext not in allowed_exts:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type. Allowed: {', '.join(allowed_exts)}"
            )
        
        file_path = None
        try:
            # Save to temp file
            file_path = f"temp_transcript_{uuid.uuid4()}{file_ext}"
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(transcript.file, buffer)
            
            # Extract text
            if file_ext == ".pdf":
                try:
                    transcript_text = DocumentProcessor.extract_text(file_path)
                    if not transcript_text or len(transcript_text.strip()) < 50:
                        raise HTTPException(
                            status_code=400,
                            detail="Transcript appears to be empty or too short. Please upload a valid transcript with at least 50 characters. If your PDF is image-based, please convert it to text first."
                        )
                except Exception as e:
                    if isinstance(e, HTTPException):
                        raise
                    raise HTTPException(
                        status_code=400,
                        detail=f"Failed to extract text from transcript: {str(e)}. If your PDF is image-based, please convert it to text first."
                    )
            else:
                # For images, we'd need OCR - for now, return error
                raise HTTPException(
                    status_code=400,
                    detail="Image transcripts require OCR. Please upload a PDF transcript."
                )
            
            # Extract courses from transcript
            if transcript_text and course_advisor_service:
                try:
                    completed_courses = course_advisor_service.extract_courses_from_text(transcript_text)
                    print(f"✅ Extracted {len(completed_courses)} courses from transcript: {completed_courses}")
                except Exception as e:
                    print(f"⚠️ Warning: Failed to extract courses from transcript: {e}")
                    completed_courses = []
            else:
                completed_courses = []
            
        finally:
            # Cleanup
            if file_path and os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except:
                    pass
        
        # Generate analysis
        print(f"🎯 Starting career analysis for: {area_of_interest}")
        print(f"📊 Transcript length: {len(transcript_text) if transcript_text else 0} chars")
        print(f"📚 Completed courses: {len(completed_courses)}")
        
        result = await academic_advisor_service.analyze_career_path(
            transcript_text=transcript_text,
            area_of_interest=area_of_interest,
            completed_courses=completed_courses
        )
        
        print(f"✅ Career analysis completed successfully")
        
        # Track event
        analytics_service.track_event(
            user="anonymous",  # Can be enhanced with session_id
            action="TRANSCRIPT_ADVISOR",
            meta={"type": "career_analysis", "area_of_interest": area_of_interest}
        )
        
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to analyze career path: {str(e)}"
        )


@app.post("/academic-advisor/chat")
async def chat_with_advisor(
    request: dict = Body(...)
):
    """
    Chat with career advisor.
    
    Request body:
    {
        "message": "user question",
        "context": {...},  # optional career analysis context
        "area_of_interest": "..."  # optional
    }
    
    Returns:
        Advisor response
    """
    if not academic_advisor_service:
        raise HTTPException(
            status_code=500,
            detail="Academic Advisor Service not initialized"
        )
    
    try:
        message = request.get("message", "")
        context = request.get("context")
        area_of_interest = request.get("area_of_interest")
        
        if not message:
            raise HTTPException(
                status_code=400,
                detail="Message is required"
            )
        
        result = await academic_advisor_service.chat_with_advisor(
            message=message,
            context=context,
            area_of_interest=area_of_interest
        )
        
        return result
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to get advisor response: {str(e)}"
        )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8008)


